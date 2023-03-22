# 从 http://www.godcom.net/ 获取圣经文本
from typing import List, Dict
import requests
from bs4 import BeautifulSoup, ResultSet
from docx import Document


# def get_version_index_url(version: str = 'niv') -> str:
#     print('开始生成版本链接')
#     return f'http://www.godcom.net/{version}/index.htm'
#
#
# def get_Bible_chapters(index_url: str = 'http://www.godcom.net/niv/index.htm') -> List[Dict]:
#     print('正在获取圣经文本结构')
#     # 文本结构  Bible_content_structure=[
#     # {'chapter_title':'Genesis',
#     # 'paragraph_links_list':['http://www.godcom.net/niv/B01C001.htm','http://www.godcom.net/niv/B01C002.htm']},
#     # {'chapter_title':'Genesis',
#     # 'paragraph_links_list':['http://www.godcom.net/niv/B01C001.htm','http://www.godcom.net/niv/B01C002.htm']},
#     # ]
#     Bible_content_structure: List[Dict] = []
#     r: str = requests.get(index_url).text
#     soup: BeautifulSoup = BeautifulSoup(r, features="lxml")
#     # 每个tr包含圣经一章和若干个小节
#     all_tr_tag: ResultSet = soup.find_all('tr')
#     for tr_tag in all_tr_tag:
#         one_chapter_structure: dict = {}
#         all_td_in_tr_tag: ResultSet = tr_tag.find_all('td')
#         chapter_title_td: ResultSet = all_td_in_tr_tag[0]
#         paragraph_links_td: ResultSet = all_td_in_tr_tag[1]
#         chapter_title: str = chapter_title_td.p.text
#         # print(chapter_title)
#         one_chapter_structure['chapter_title'] = chapter_title
#         one_chapter_paragraph_links_list: list = []
#         all_a_in_tr_tag: ResultSet = paragraph_links_td.find_all('a')
#         for link in all_a_in_tr_tag:
#             paragraph_link: str = 'http://www.godcom.net/niv/' + link.get('href')
#             one_chapter_paragraph_links_list.append(paragraph_link)
#             # print(link.get('href'))
#         one_chapter_structure['paragraph_links_list'] = one_chapter_paragraph_links_list
#         Bible_content_structure.append(one_chapter_structure)
#     # print('获取圣经文本结构完成')
#     return Bible_content_structure
def get_Bible_chapters(index_url: str = 'http://www.godcom.net/niv/index.htm') -> List[Dict]:
    print('正在获取圣经文本结构')
    # 文本结构
    # Bible_content_structure = [
    #     {'chapter_title': 'Genesis', 'paragraph_links_list': [{'1': 'http://www.godcom.net/niv/B01C001.htm'},
    #                                                           {'2': 'http://www.godcom.net/niv/B01C002.htm'}]},
    #     {'chapter_title': 'Genesis', 'paragraph_links_list': [{'1': 'http://www.godcom.net/niv/B01C001.htm'},
    #                                                           {'2': 'http://www.godcom.net/niv/B01C002.htm'}]},
    # ]
    Bible_content_structure: List[Dict] = []
    r: str = requests.get(index_url).text
    soup: BeautifulSoup = BeautifulSoup(r, features="lxml")
    # 每个tr包含圣经一章和若干个小节,结构为里面包含2个td，第一个是章节名称，第二个是小节序号和链接
    all_tr_tag: ResultSet = soup.find_all('tr')
    for tr_item in all_tr_tag:
        # 生命单一章节数据结构
        one_chapter_dict: dict = {}
        paragraph_links_list: List[Dict[str, str]] = []
        # one_paragraph_dict: Dict[str, str] = {}
        # 获取单一章节数据
        all_td_tag_in_tr_item: ResultSet = tr_item.find_all('td')
        # 获取每章标题
        chapter_name_td_tag: ResultSet = all_td_tag_in_tr_item[0]
        chapter_title: str = chapter_name_td_tag.text
        one_chapter_dict['chapter_title'] = chapter_title
        # 包含所有小节a标签的td列表
        all_paragraph_td_tag: ResultSet = all_td_tag_in_tr_item[1]
        for paragraph_td_tag in all_paragraph_td_tag.find_all('a'):
            one_paragraph_dict: Dict[str, str] = {}
            one_paragraph_order: str = paragraph_td_tag.text
            one_paragraph_link: str = 'http://www.godcom.net/niv/' + paragraph_td_tag.get('href')
            # print(one_paragraph_link)
            one_paragraph_dict['order'] = one_paragraph_order
            one_paragraph_dict['link'] = one_paragraph_link
            paragraph_links_list.append(one_paragraph_dict)
        one_chapter_dict['paragraph_links_list'] = paragraph_links_list
        Bible_content_structure.append(one_chapter_dict)

    return Bible_content_structure


def write_content_to_docx(file_name: str = 'NIV Bible', path: str = '.', file_type: str = 'docx') -> None:
    document = Document()
    path_and_name: str = path + '/' + file_name + '.' + file_type
    Bible_content_structure: List[Dict] = get_Bible_chapters()
    for item in Bible_content_structure:
        chapter_title: str = item['chapter_title']
        document.add_heading(chapter_title)
        print(f'{chapter_title}章开始写入')
        for paragraph in item['paragraph_links_list']:
            paragraph_order = paragraph['order']
            paragraph_link = paragraph['link']
            paragraph_text_list: List[str] = get_paragraph_content(paragraph_link)
            document.add_heading(paragraph_order, level=2)
            for text in paragraph_text_list:
                document.add_paragraph(text)
            print(f'{chapter_title}第{paragraph_order}小节写入完成')
        # document.add_break()
    document.save(path_and_name)
    print('=====================全部文档写入成功=====================')


def get_paragraph_content(paragraph_url: str = 'http://www.godcom.net/niv/B01C001.htm') -> List[str]:
    paragraph_text_list: List[str] = []
    r = requests.get(paragraph_url).text
    soup = BeautifulSoup(r, features="lxml")
    try:
        whole_paragraph_text_contain_num: ResultSet = soup.find_all('table')[1]
    except IndexError:
        whole_paragraph_text_contain_num: ResultSet = soup.find_all('table')[0]
    tr_in_whole_paragraph_text_contain_num: ResultSet = whole_paragraph_text_contain_num.find_all('tr')[:-1]
    for item in tr_in_whole_paragraph_text_contain_num:
        text_td = item.find_all('td')[1]
        paragraph_text_list.append(text_td.text)
    # print(whole_paragraph_text_contain_num.find_all('tr'))
    return paragraph_text_list


if __name__ == '__main__':
    write_content_to_docx()
    # print(get_paragraph_content())
