from docx import Document

title_fist='这是个一级标题'
title_second='这是个二级标题'

text1='这里是正文1这里是正文1这里是正文1'
text2='这里是正文2这里是正文2这里是正文2\n这里是正文2这里是正文2这里是正文2'

text_list=['1','2','3','4','5']

document = Document()
# document.add_paragraph(text1)
# document.add_heading(title_fist)
# document.add_paragraph(text2)
# document.add_heading(title_second,level=2)

for item in text_list:
      document.add_paragraph(item)

document.save(r"./test.docx")